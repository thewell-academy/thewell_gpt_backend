import json
import uuid
from typing import Dict, List
import os

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from sqlalchemy import text
from sqlalchemy.orm import Session
from starlette import status
from database.models.answer_option_info import AnswerOptionInfo
from database.models.default_question_info import DefaultQuestionInfo
from database.models.exam_question import ExamQuestion
from database.models.subject_detail import SubjectDetail
from database.pydantic_models.pydantic_models import ExamQuestionCreate, QuestionRequest
from sqlalchemy.orm.attributes import flag_modified
import uuid

from service.question_bank.question_bank_util import TableFlowManager, get_passage_text, get_passage_text


async def save_exam_question(question_request: QuestionRequest, replace: bool, db: Session):

    exam_question_data = question_request.question_model

    try:
        subject = exam_question_data.subject
        grade = exam_question_data.default_question_info.grade
        exam = exam_question_data.default_question_info.exam
        exam_year = exam_question_data.default_question_info.exam_year
        exam_month = exam_question_data.default_question_info.exam_month

        question_numbers_list = [str(i.question_number) for i in exam_question_data.answer_option_info_list]
        question_numbers = ",".join(question_numbers_list) if question_numbers_list else ""

        exist = same_question_exists(exam, exam_year, exam_month, question_numbers, subject, grade, db)

        if exist and replace:
            existing_question = db.query(ExamQuestion).filter(ExamQuestion.id == exist).first()
            if existing_question:
                existing_question.valid = False
                db.commit()
                db.refresh(existing_question)
            else:
                return {
                    "status_code": status.HTTP_200_OK,
                    "detail": "Existing question not found."
                }

        if exist and not replace:
            return {
                "status_code": status.HTTP_203_NON_AUTHORITATIVE_INFORMATION,
            }

        db_default_info = DefaultQuestionInfo(
            exam=exam_question_data.default_question_info.exam,
            exam_year=exam_question_data.default_question_info.exam_year,
            exam_month=exam_question_data.default_question_info.exam_month,
            grade=exam_question_data.default_question_info.grade,
            file_path=exam_question_data.default_question_info.file_path,
            selected_file_bytes=exam_question_data.default_question_info.selected_file_bytes
        )

        db_exam_question = ExamQuestion(
            subject=exam_question_data.subject,
            type=question_request.question_type,
            question_content_text_map=exam_question_data.question_content_text_map,
            question_numbers=question_numbers,  # Now defined in the model
            default_question_info=db_default_info  # Correct usage
        )

        for answer_option_data in exam_question_data.answer_option_info_list:
            db_answer_option = AnswerOptionInfo(
                question_number=answer_option_data.question_number,
                question_score=answer_option_data.question_score,
                abc_option_list=answer_option_data.abc_option_list,
                question_text=answer_option_data.question_text,
                option1=answer_option_data.options[0],
                option2=answer_option_data.options[1],
                option3=answer_option_data.options[2],
                option4=answer_option_data.options[3],
                option5=answer_option_data.options[4],
                answer=answer_option_data.selected_answer,
                memo=answer_option_data.memo
            )
            db_exam_question.answer_option_info_list.append(db_answer_option)

        db.add(db_exam_question)
        db.commit()
        db.refresh(db_exam_question)
        return {
            "status_code": status.HTTP_200_OK,
        }
    except Exception as e:
        return {
            "status_code": status.HTTP_500_INTERNAL_SERVER_ERROR
        }


def same_question_exists(exam, exam_year, exam_month, question_numbers, subject, grade, db: Session):
    if not isinstance(exam_year, int):
        exam_year = int(exam_year)
    if not isinstance(exam_month, int):
        exam_month = int(exam_month)
    if not isinstance(question_numbers, str):
        question_numbers = str(question_numbers)

    existing_question = db.query(ExamQuestion).join(DefaultQuestionInfo).filter(
        ExamQuestion.subject == subject,
        ExamQuestion.valid == True,
        ExamQuestion.question_numbers == question_numbers,
        DefaultQuestionInfo.exam == exam,
        DefaultQuestionInfo.exam_year == exam_year,
        DefaultQuestionInfo.exam_month == exam_month,
        DefaultQuestionInfo.grade == grade,
    ).first()

    if existing_question:
        return existing_question.id
    else:
        return None


async def save_subject_details(subject: str, details: dict, db: Session, parent_id: int = None, path: str = ""):
    for key, value in details.items():
        current_path = f"{path} > {key}" if path else key

        if isinstance(value, list):
            value = {}

        existing_node = db.query(SubjectDetail).filter_by(subject=subject, name=key, parent_id=parent_id).first()

        if existing_node:
            if isinstance(value, dict):
                await save_subject_details(subject, value, db, existing_node.id, current_path)
            else:
                existing_node.value = value
                db.commit()
        else:
            new_node = SubjectDetail(
                subject=subject,
                name=key,
                parent_id=parent_id,
                path=current_path,
                value=None if isinstance(value, dict) else value,
            )
            db.add(new_node)
            db.commit()
            db.refresh(new_node)

            if isinstance(value, dict):
                await save_subject_details(subject, value, db, new_node.id, current_path)


async def get_subject_details_data(subject: str, db: Session):
    nodes = db.query(SubjectDetail).filter_by(subject=subject).all()

    node_map = {node.id: node for node in nodes}

    hierarchy = {}

    for node in nodes:
        current_entry = {}

        if node.parent_id:
            parent = node_map[node.parent_id]
            parent_entry = hierarchy
            for part in parent.path.split(" > "):
                parent_entry = parent_entry.setdefault(part, {})
            parent_entry[node.name] = current_entry
        else:
            hierarchy[node.name] = current_entry

    return hierarchy


async def delete_question(question_id, db: Session):
    exam_question = db.query(ExamQuestion).filter(ExamQuestion.id == question_id).first()

    if not exam_question:
        return {"message": "ExamQuestion not found"}

    default_question_info_id = exam_question.default_question_info_id

    db.delete(exam_question)
    db.commit()

    if default_question_info_id:
        is_referenced = (
            db.query(ExamQuestion)
            .filter(ExamQuestion.default_question_info_id == default_question_info_id)
            .count()
        )

        if is_referenced == 0:
            default_question_info = db.query(DefaultQuestionInfo).filter(
                DefaultQuestionInfo.id == default_question_info_id).first()
            if default_question_info:
                db.delete(default_question_info)
                db.commit()

    return {"message": "ExamQuestion and all related data deleted successfully"}

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.section import WD_SECTION_START


def is_page_blank(paragraphs, tables):
    """
    Check if a collection of paragraphs and tables on a page is blank.
    A page is considered blank if all paragraphs and tables only contain whitespace or are empty.
    """
    # Check paragraphs
    for paragraph in paragraphs:
        if paragraph.text.strip():  # If there's visible text, the page is not blank
            return False

    # Check tables
    for table in tables:
        for row in table.rows:
            for cell in row.cells:
                if any(paragraph.text.strip() for paragraph in cell.paragraphs):  # Check cell content
                    return False

    return True


def remove_blank_pages(doc):
    """
    Removes blank pages from a Word document by identifying and excluding sections with no content.
    """
    # Create a new document to store non-blank content
    new_doc = Document()

    # Remove the default blank paragraph in the new document
    if new_doc.paragraphs:
        p = new_doc.paragraphs[0]
        p._element.getparent().remove(p._element)

    # Iterate over all paragraphs in the document
    paragraphs = doc.paragraphs
    tables = doc.tables

    # Group content by sections/pages
    current_page_content = {"paragraphs": [], "tables": []}
    all_pages_content = []

    for paragraph in paragraphs:
        if paragraph.text.startswith("\f"):  # Page break indicator
            # Store the current page content and start a new one
            all_pages_content.append(current_page_content)
            current_page_content = {"paragraphs": [], "tables": []}
        else:
            current_page_content["paragraphs"].append(paragraph)

    # Append the last page's content
    all_pages_content.append(current_page_content)

    # Remove blank pages
    for page_content in all_pages_content:
        if not is_page_blank(page_content["paragraphs"], tables):
            for paragraph in page_content["paragraphs"]:
                new_p = new_doc.add_paragraph(paragraph.text)
                new_p.style = paragraph.style

    return new_doc


async def export_question_service(
        subject: str,
        exam: str,
        selections: List[str],
        years: List[int],
        months: List[int],
        grades: List[str],
        db: Session
):
    if exam == "수능":
        existing_question_query_result = db.query(ExamQuestion).join(DefaultQuestionInfo).filter(
            ExamQuestion.valid == True,
            ExamQuestion.subject == subject,
            DefaultQuestionInfo.exam == exam,
            ExamQuestion.type.in_(selections),
            DefaultQuestionInfo.exam_year.in_(years),
        ).all()
    else:
        existing_question_query_result = db.query(ExamQuestion).join(DefaultQuestionInfo).filter(
            ExamQuestion.valid == True,
            ExamQuestion.subject == subject,
            DefaultQuestionInfo.exam == exam,
            ExamQuestion.type.in_(selections),
            DefaultQuestionInfo.exam_year.in_(years),
            DefaultQuestionInfo.exam_month.in_(months),
            DefaultQuestionInfo.grade.in_(grades),
        ).all()

    existing_question_list: List[ExamQuestion] = [
        question() if isinstance(question, type) else question
        for question in existing_question_query_result
    ]

    doc = Document()

    section = doc.sections[0]
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

    default_font_path = "default_font.ttf"  # Replace with your font file path
    if not os.path.exists(default_font_path):
        raise FileNotFoundError(f"Font file '{default_font_path}' not found in the working directory.")

    # Update the Normal style
    style = doc.styles['Normal']
    font = style.font
    font.size = Pt(8)  # Set font size to 9pt

    # Set custom font using font file
    font.name = 'CustomFont'  # Logical font name
    font.element.rPr.rFonts.set(qn('w:ascii'), 'CustomFont')  # Applies to ASCII text
    font.element.rPr.rFonts.set(qn('w:eastAsia'), 'CustomFont')  # Applies to East Asian text
    font.element.rPr.rFonts.set(qn('w:hAnsi'), 'CustomFont')  # Applies to high ANSI text
    font.element.rPr.rFonts.set(qn('w:cs'), 'CustomFont')  # Applies to complex scripts

    manager = TableFlowManager(
        doc,
        max_lines_per_cell=25,
        max_chars_per_line=70,
    )

    answer_list = []
    for i, exam_question_class in enumerate(existing_question_list):
        passage_text = get_passage_text(exam_question_class)

        manager.add_question(
            passage_text=passage_text,
            subquestion_list=[
                (
                    i.question_text,
                    [i.option1, i.option2, i.option3, i.option4, i.option5]
                )
                for i in exam_question_class.answer_option_info_list
            ],
            file_bytes=exam_question_class.default_question_info.selected_file_bytes
        )

        for k in exam_question_class.answer_option_info_list:
            answer_list.append(k.answer)

    manager.add_answers([
        (i + 1, answer) for i, answer in enumerate(answer_list)
    ])

    output_file = f"{str(uuid.uuid4())}.docx"
    doc.save(output_file)

    return output_file
