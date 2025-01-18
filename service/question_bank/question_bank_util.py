from typing import List, Dict, Tuple
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import textwrap
from io import BytesIO
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

from docx.table import _Cell

from lxml import etree
import re
from latex2mathml.converter import convert
import mathml2omml

from database.models.exam_question import ExamQuestion

font_name = "Times New Roman"
font_size = 9


def get_passage_text(exam_question: ExamQuestion):
    if exam_question.subject != "영어":
        return ""

    question_type_normal = ["글의 목적", "글의 분위기 / 심경", "대의 파악", "함의 추론", "도표 이해",
                            "내용 일치 / 불일치", "실용문 일치 / 불일치", "어법성 판단",
                            "단어 쓰임 판단", "빈칸 추론", "무관한 문장", "주어진 문장 넣기"]
    question_type_order = ["글의 순서"]
    question_type_summary = ["요약문 완성"]
    question_type_long1 = ["기본 장문 독해"]
    question_type_long2 = ["복합 문단 독해"]

    question_content_text_map: Dict[str] = exam_question.question_content_text_map
    keys = list(question_content_text_map.keys())
    if exam_question.type in question_type_normal:
        big_text = question_content_text_map[keys[0]]
    elif exam_question.type in question_type_order:
        big_text = f"""
        {question_content_text_map[keys[0]]}\n\n
            
        (A) {question_content_text_map[keys[1]]}\n\n
        
        (B) {question_content_text_map[keys[2]]}\n\n
        
        (C) {question_content_text_map[keys[3]]}
        """
    elif exam_question.type in question_type_summary:
        big_text = f"""
        {question_content_text_map[keys[0]]}\n\n
        
        {question_content_text_map[keys[1]]}
        """
    elif exam_question.type in question_type_long1:
        big_text = question_content_text_map[keys[0]]
    elif exam_question.type in question_type_long2:
        big_text = f"""
        {question_content_text_map[keys[0]]}\n\n

        (A) {question_content_text_map[keys[1]]}\n\n

        (B) {question_content_text_map[keys[2]]}\n\n

        (C) {question_content_text_map[keys[3]]}
        """
    else:
        big_text = ""

    big_text = big_text.strip()
    big_text = big_text.replace("\n", "")

    return big_text


#
# ---------- HELPER FUNCTIONS ----------
#
def set_justification(paragraph, alignment="both"):
    p = paragraph._element
    pPr = p.get_or_add_pPr()
    jc = OxmlElement("w:jc")
    jc.set(qn("w:val"), alignment)
    pPr.append(jc)


def add_formatted_text(paragraph, text):
    parts = re.split(r"(<\/?[biu]>)", text)
    formatting = {"b": False, "i": False, "u": False}

    for part in parts:
        if part in ("<b>", "</b>"):
            formatting["b"] = (part == "<b>")
        elif part in ("<i>", "</i>"):
            formatting["i"] = (part == "<i>")
        elif part in ("<u>", "</u>"):
            formatting["u"] = (part == "<u>")
        else:
            run = paragraph.add_run(part)
            run.font.bold = formatting["b"]
            run.font.italic = formatting["i"]
            run.font.underline = formatting["u"]
            run.font.name = font_name
            run.font.size = Pt(font_size)


def add_paragraph_with_alignment(cell: _Cell, text, alignment="both"):
    print(text.splitlines()[0])
    if len(text.splitlines()[0]) < 30 or re.match(r"^\s*\(?\d+\)", text):
        alignment = "left"
        print("[[align left]]")
    else:
        alignment = "both"
        print("[[align both]]")

    if cell.paragraphs and cell.paragraphs[0].text.strip() == "":
        p = cell.paragraphs[0]
    else:
        p = cell.add_paragraph()

    add_formatted_text(p, text)
    set_justification(p, alignment)
    return p


def split_text_by_lines(text: str,
                        max_lines: int = 20,
                        max_chars_per_line: int = 70) -> List[str]:
    raw_lines = []
    paragraphs = text.split("\n")
    for para in paragraphs:
        wrapped = textwrap.wrap(para, width=max_chars_per_line)
        if not wrapped:
            raw_lines.append("")
        else:
            raw_lines.extend(wrapped)

    chunks = []
    current_chunk = []
    for line in raw_lines:
        current_chunk.append(line)
        if len(current_chunk) >= max_lines:
            chunks.append("\n".join(current_chunk))
            current_chunk = []
    # leftover
    if current_chunk:
        chunks.append("\n".join(current_chunk))

    return chunks


def create_omml_element(latex_code):
    try:
        mathml = convert(latex_code)
    except Exception as e:
        raise ValueError(f"Error converting LaTeX to MathML: {e}")

    try:
        omml = mathml2omml.convert(mathml)
    except Exception as e:
        raise ValueError(f"Error converting MathML to OMML: {e}")

    if 'xmlns:m=' not in omml:
        omml = omml.replace(
            "<m:oMath",
            '<m:oMath xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math"'
        )

    try:
        omml_element = etree.fromstring(omml.encode('utf-8'))
    except Exception as e:
        raise ValueError(f"Error parsing OMML: {e}")

    return omml_element


def insert_omml(paragraph, latex_code):
    omml_element = create_omml_element(latex_code)
    paragraph._element.append(omml_element)


class TableFlowManager:
    def __init__(self,
                 doc: Document,
                 max_lines_per_cell=20,
                 max_chars_per_line=70):
        self.doc = doc
        self.max_lines_per_cell = max_lines_per_cell
        self.max_chars_per_line = max_chars_per_line

        self.current_table = self._create_new_table()
        self.cell_order = [(0, 0), (1, 0), (0, 1), (1, 1)]
        self.cell_index = 0
        self.page_size = 1
        self.question_number = 0

    def _create_new_table(self):
        table = self.doc.add_table(rows=2, cols=2)
        table.autofit = False

        table.columns[0].width = Inches(3.75)
        table.columns[1].width = Inches(3.75)

        row_height_twips = 7150  # 12.62 cm in Twips
        for row in table.rows:
            tr = row._tr
            trHeight = OxmlElement('w:trHeight')
            trHeight.set(qn('w:val'), str(row_height_twips))
            trHeight.set(qn('w:hRule'), 'exact')
            tr.append(trHeight)

        return table

    def _start_new_page_and_table(self):
        if self.cell_index >= 4:
            self.current_table = self._create_new_table()
            self.cell_index = 0

    def _get_next_cell(self):
        if self.cell_index >= 4:
            self._start_new_page_and_table()

        r, c = self.cell_order[self.cell_index]
        self.cell_index += 1
        return r, c

    def _get_next_question_number(self):
        self.question_number += 1
        return self.question_number

    def add_question(
            self,
            passage_text: str,
            subquestion_list: List[Tuple[str, List[str]]] = None,
            file_bytes: bytes = None
    ):
        if subquestion_list is None:
            subquestions = []

        for subquestion in subquestion_list:
            question_text_list: List[Dict[str, str]] = eval(
                subquestion[0].replace("true", "True").replace("false", "False")
            )
            question_text_list.insert(0, {"insert": f"{self._get_next_question_number()}. "})

            answer_options_list_is_empty = not any(subquestion[1])
            answer_options_list: List[List[Dict[str, str]]] = [
                eval(i.replace("true", "True").replace("false", "False"))
                for i in subquestion[1]
            ] if not answer_options_list_is_empty else []

            self.add_question_to_cell(
                question_text_list,
                answer_options_list,
                file_bytes
            )

    def add_question_to_cell(self,
                             question_text_list: List[Dict[str, str]],
                             answer_option_list: List[List[Dict[str, str]]],
                             file_bytes: bytes = None):

        def rearrange_text_list(text_list: List[Dict[str, str]]):
            if not text_list:  # Handle empty input
                return []

            result = []

            for item in text_list:
                insert_value = item.get("insert", "")
                if insert_value.strip() == "":  # If the insert is empty or only contains \n
                    if result:
                        # Append the empty or newline-only value to the previous item's insert
                        result[-1]["insert"] += insert_value
                else:
                    # Add non-empty items to the result list
                    result.append(item)

            return result

        def process_text_parts(paragraph, content_text, content_attributes, process_cell):
            """
            Processes text parts and handles attributes like bold, italic, underline, and box.
            If 'box' is True, creates a box/table for the text directly and resizes it.
            """
            parts = re.split(latex_pattern, content_text)
            for idx, part_text in enumerate(parts):
                if idx % 2 == 0:  # Regular text
                    if content_attributes.get("box", False) == True:  # Handle 'box' attribute
                        # Create a 1x1 table inside the cell for the box
                        box_table = process_cell.add_table(rows=1, cols=1)
                        box_table.style = "Table Grid"

                        # Set the table width to fit within the cell
                        box_table.autofit = False
                        parent_cell_width = process_cell.width if process_cell.width else Inches(1)

                        table_width = parent_cell_width * 0.5
                        for row in box_table.rows:
                            for box_cell in row.cells:
                                box_cell.width = table_width  # Set the cell width

                        # Access the underlying XML of the table
                        tbl = box_table._tbl

                        # Add tblPr if it doesn't exist
                        tbl_pr = tbl.find(qn('w:tblPr'))
                        if tbl_pr is None:
                            tbl_pr = OxmlElement('w:tblPr')
                            tbl.insert(0, tbl_pr)

                        # Add tblBorders to tblPr
                        tbl_borders = OxmlElement('w:tblBorders')

                        for border_name in ['top', 'left', 'bottom', 'right']:
                            border = OxmlElement(f'w:{border_name}')
                            border.set(qn('w:val'), 'single')  # Border style
                            border.set(qn('w:sz'), '4')  # Border thickness
                            border.set(qn('w:space'), '0')  # Border spacing
                            border.set(qn('w:color'), '000000')  # Border color
                            tbl_borders.append(border)

                        tbl_pr.append(tbl_borders)

                        # Add text to the box
                        box_cell = box_table.cell(0, 0)
                        box_cell.width = table_width
                        box_paragraph = box_cell.paragraphs[0]
                        run = box_paragraph.add_run(part_text.rstrip("\n"))
                        run.font.name = font_name
                        run.font.size = Pt(font_size)
                        if content_attributes.get("bold"):
                            run.bold = True
                        if content_attributes.get("underline"):
                            run.underline = True
                        if content_attributes.get("italic"):
                            run.italic = True

                        box_paragraph.add_run("\n")  # Add a line break for box content
                    else:
                        # Handle regular text without 'box'
                        run = paragraph.add_run(part_text.rstrip("\n"))
                        run.font.name = font_name
                        run.font.size = Pt(font_size)
                        if content_attributes.get("bold"):
                            run.bold = True
                        if content_attributes.get("underline"):
                            run.underline = True
                        if content_attributes.get("italic"):
                            run.italic = True
                else:  # LaTeX content
                    insert_omml(paragraph, part_text)

        def add_box_to_cell(cell, box_content_list):
            """
            Adds a single table to the cell with grouped box content.
            """
            box_table = cell.add_table(rows=1, cols=1)
            box_table.style = "Table Grid"
            box_cell = box_table.cell(0, 0)

            box_table.autofit = False
            parent_cell_width = cell.width if cell.width else Inches(3)  # Default to 3 inches if width is not set
            table_width = parent_cell_width * 0.95  # Slightly less than the parent width for margin
            for row in box_table.rows:
                for box_cell in row.cells:
                    box_cell.width = table_width

            # Set black border styling for the table
            tbl = box_table._tbl
            tbl_pr = tbl.find(qn('w:tblPr'))
            if tbl_pr is None:
                tbl_pr = OxmlElement('w:tblPr')
                tbl.insert(0, tbl_pr)

            tbl_borders = parse_xml(
                r'<w:tblBorders %s>'
                r'<w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
                r'<w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
                r'<w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
                r'<w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
                r'</w:tblBorders>' % nsdecls('w')
            )
            tbl_pr.append(tbl_borders)

            # Add each line into the table
            box_paragraph = box_cell.paragraphs[0]
            for box_text, box_attributes in box_content_list:
                process_text_parts(box_paragraph, box_text, box_attributes, box_cell)
                box_paragraph.add_run("\n")  # Add a line break for each box line

        question_text_list = rearrange_text_list(question_text_list)
        for i in range(len(answer_option_list)):
            answer_option_list[i] = rearrange_text_list(answer_option_list[i])

        r, c = self._get_next_cell()
        cell = self.current_table.cell(r, c)
        latex_pattern = r"\[:(.*?)\]"

        box_group = []  # Holds grouped "box" content

        for question_text in question_text_list:
            text_content = question_text.get("insert", "")
            text_attributes = question_text.get("attributes", {})
            is_box = text_attributes.get("box", False)

            if is_box:
                # Collect content for the box
                box_group.append((text_content, text_attributes))
            else:
                # If exiting box mode, add the collected box content as a table
                if box_group:
                    add_box_to_cell(cell, box_group)
                    box_group.clear()  # Clear the group after processing
                    cell.add_paragraph("\n")  # Add a line break after the box

                # Add regular text
                cell_paragraph = cell.paragraphs[0]
                process_text_parts(cell_paragraph, text_content, text_attributes, cell)

        # If any box content remains unprocessed, add it as a table
        if box_group:
            add_box_to_cell(cell, box_group)

        cell.paragraphs[0].add_run("\n\n")
        # Add the image, if provided
        if file_bytes is not None:
            image_stream = BytesIO(file_bytes)
            run = cell.paragraphs[0].add_run()
            picture = run.add_picture(image_stream)

            # Resize the image to fit the cell width
            cell_width = int(cell.width / 2) if hasattr(cell, "width") else Inches(2)
            image_width = picture.width
            image_height = picture.height

            # Calculate new height maintaining aspect ratio
            new_width = cell_width
            aspect_ratio = image_height / image_width
            new_height = int(new_width * aspect_ratio)
            picture.width = new_width
            picture.height = new_height
            cell.paragraphs[0].add_run("\n\n")

        # Add answer options
        for num, answer_option in enumerate(answer_option_list):
            for option_item in answer_option:
                option_text = f"({num + 1}) " + option_item.get("insert", "")
                option_text = option_text.replace("\n", "")
                option_attributes = option_item.get("attributes", {})
                process_text_parts(cell.paragraphs[0], option_text, option_attributes, cell)
                cell.paragraphs[0].add_run("\n")

        if r == 1 and c == 1:
            self._start_new_page_and_table()

    def add_answers(self, answer_list: List[tuple[int, int]]):
        self.doc.add_page_break()

        # Add a heading for the answers section
        heading = self.doc.add_paragraph()
        heading.add_run("정답").bold = True
        heading.style = "Heading 1"

        # Loop through the answer list and write each answer
        for question_number, answer in answer_list:
            # Add a new paragraph for each answer
            answer_paragraph = self.doc.add_paragraph()
            answer_text = f"({question_number}) {answer}"
            run = answer_paragraph.add_run(answer_text)

            # Set font styles for the text
            run.font.name = font_name
            run.font.size = Pt(font_size)
